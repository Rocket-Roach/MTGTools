#!/usr/bin/env python3
"""
MTG Modern Gateway Deck Tracker
Tracks card collection progress towards building Esper Blink gateway deck and upgrades.
"""

import json
import os
import sys
import re
import csv
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, asdict
from collections import defaultdict

# Optional dependencies - handle gracefully
try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pyperclip
    HAS_CLIPBOARD = True
except ImportError:
    HAS_CLIPBOARD = False


@dataclass
class Card:
    name: str
    quantity: int = 1
    set_code: str = ""
    collector_number: str = ""
    foil: bool = False

    def key(self) -> str:
        # Printing-agnostic: same card name from any set is the same card.
        return self.name.strip().lower()

    def display_name(self) -> str:
        # Printing-agnostic: quantity + name only.
        return f"{self.quantity}x {self.name}"


def card_faces(name: str) -> set:
    """Normalized face names for matching (splits modal double-faced '//' names).

    "Boggart Trawler // Boggart Bog" -> {"boggart trawler", "boggart bog"}.
    Single-faced names return a one-element set. Either side of a '//' card
    matches a requirement for either side (or the full name).
    """
    faces = set()
    for part in re.split(r'\s*/+\s*', (name or '').strip()):
        part = re.sub(r'\s+', ' ', part).strip().lower()
        if part:
            faces.add(part)
    return faces or {''}


@dataclass
class Collection:
    cards: Dict[str, Card]

    def __init__(self):
        self.cards = {}
        # face -> set of entry keys, and face -> total quantity.
        # Keeps add()/get_quantity() near-O(1) instead of scanning + regex
        # over every entry on every lookup.
        self._face_index: Dict[str, set] = {}
        self._face_qty: Dict[str, int] = {}

    def _index_add(self, key: str, name: str, qty: int):
        for f in card_faces(name):
            s = self._face_index.get(f)
            if s is None:
                self._face_index[f] = {key}
            else:
                s.add(key)
            self._face_qty[f] = self._face_qty.get(f, 0) + qty

    def _index_remove(self, key: str, name: str, qty: int):
        for f in card_faces(name):
            s = self._face_index.get(f)
            if s is not None:
                s.discard(key)
                if not s:
                    del self._face_index[f]
            left = self._face_qty.get(f, 0) - qty
            if left > 0:
                self._face_qty[f] = left
            else:
                self._face_qty.pop(f, None)

    def add(self, card: Card):
        key = card.key()
        if key in self.cards:
            self.cards[key].quantity += card.quantity
            self._index_add(key, self.cards[key].name, card.quantity)
            return
        # Merge with an entry sharing any face, so "Wear" and
        # "Wear // Tear" never become duplicate rows (prefer full name).
        cand = set()
        for f in card_faces(card.name):
            cand |= self._face_index.get(f, set())
        cand.discard(key)
        if cand:
            ek = next(iter(cand))
            existing = self.cards[ek]
            existing.quantity += card.quantity
            self._index_add(ek, existing.name, card.quantity)
            if '//' in card.name and '//' not in existing.name:
                self._index_remove(ek, existing.name, existing.quantity)
                existing.name = card.name
                self._index_add(ek, existing.name, existing.quantity)
            return
        self.cards[key] = card
        self._index_add(key, card.name, card.quantity)

    def remove(self, card: Card, quantity: int = 1):
        key = card.key()
        if key in self.cards:
            existing = self.cards[key]
            take = min(quantity, existing.quantity)
            existing.quantity -= take
            self._index_remove(key, existing.name, take)
            if existing.quantity <= 0:
                del self.cards[key]

    def pop(self, key: str):
        """Remove a whole entry by key (keeps the face index consistent)."""
        card = self.cards.pop(key, None)
        if card is not None:
            self._index_remove(key, card.name, card.quantity)
        return card

    def clear(self):
        self.cards.clear()
        self._face_index.clear()
        self._face_qty.clear()

    def get_quantity(self, name: str, set_code: str = "", collector_number: str = "", foil: bool = False) -> int:
        # Printing-agnostic AND face-aware: set/collector/foil are ignored,
        # and any entry sharing a face with the requirement counts (once).
        if not (name or '').strip():
            return 0
        seen = set()
        total = 0
        for f in card_faces(name):
            for key in self._face_index.get(f, ()):
                if key not in seen:
                    card = self.cards.get(key)
                    if card is not None:
                        seen.add(key)
                        total += card.quantity
        return total

    def total_cards(self) -> int:
        return sum(c.quantity for c in self.cards.values())

    def unique_cards(self) -> int:
        return len(self.cards)

    def to_dict(self) -> dict:
        return {k: asdict(v) for k, v in self.cards.items()}

    @classmethod
    def from_dict(cls, data: dict) -> 'Collection':
        coll = cls()
        for k, v in data.items():
            # Merge by normalized name so old printing-split files collapse
            # into single per-card entries.
            try:
                coll.add(Card(**v))
            except Exception:
                continue
        return coll


class DeckPlan:
    def __init__(self, plan_file: str):
        with open(plan_file, 'r') as f:
            self.data = json.load(f)
        self.phases = self.data['phases']
        self.summary = self.data['summary']
        self.metadata = self.data['metadata']

    def get_all_required_cards(self) -> Dict[str, List[Tuple[int, str, dict]]]:
        """Returns {card_name: [(qty, phase_name, card_info), ...]}"""
        required = defaultdict(list)
        for phase in self.phases:
            phase_name = phase['name']
            for card_info in phase['cards']:
                name = card_info['name']
                qty = card_info['quantity']
                required[name].append((qty, phase_name, card_info))
        return required

    def get_phase_cards(self, phase_name: str) -> List[dict]:
        for phase in self.phases:
            if phase['name'] == phase_name:
                return phase['cards']
        return []


class ProgressTracker:
    def __init__(self, plan: DeckPlan, collection: Collection):
        self.plan = plan
        self.collection = collection

    def check_card(self, name: str, set_code: str = "", collector_number: str = "", foil: bool = False) -> int:
        return self.collection.get_quantity(name, set_code, collector_number, foil)

    def phase_progress(self, phase_name: str) -> Tuple[int, int, float, List[dict]]:
        """Returns (owned, total, percentage, missing_cards)"""
        phase_cards = self.plan.get_phase_cards(phase_name)
        owned = 0
        total = 0
        missing = []

        for card_info in phase_cards:
            name = card_info['name']
            needed = card_info['quantity']
            have = self.check_card(name)
            owned += min(have, needed)
            total += needed
            if have < needed:
                missing.append({
                    'name': name,
                    'needed': needed,
                    'have': have,
                    'short': needed - have,
                    'price': card_info.get('price', 0),
                    'category': card_info.get('category', ''),
                    'notes': card_info.get('notes', '')
                })

        pct = (owned / total * 100) if total > 0 else 100
        return owned, total, pct, missing

    def all_phases_progress(self) -> List[dict]:
        results = []
        for phase in self.plan.phases:
            owned, total, pct, missing = self.phase_progress(phase['name'])
            results.append({
                'phase': phase['name'],
                'owned': owned,
                'total': total,
                'percentage': round(pct, 1),
                'missing_count': len(missing),
                'missing_value': sum(m['short'] * m['price'] for m in missing),
                'missing': missing
            })
        return results

    def overall_progress(self) -> dict:
        all_req = self.plan.get_all_required_cards()
        total_unique = len(all_req)
        owned_unique = 0
        total_copies_needed = 0
        total_copies_owned = 0
        total_value_missing = 0

        for name, entries in all_req.items():
            have = self.check_card(name)
            max_needed = max(qty for qty, _, _ in entries)
            total_copies_needed += max_needed
            total_copies_owned += min(have, max_needed)
            if have >= max_needed:
                owned_unique += 1
            else:
                total_value_missing += (max_needed - have) * entries[0][2].get('price', 0)

        return {
            'unique_cards_owned': owned_unique,
            'unique_cards_total': total_unique,
            'unique_percentage': round(owned_unique / total_unique * 100, 1) if total_unique > 0 else 100,
            'copies_owned': total_copies_owned,
            'copies_needed': total_copies_needed,
            'copies_percentage': round(total_copies_owned / total_copies_needed * 100, 1) if total_copies_needed > 0 else 100,
            'estimated_value_missing': total_value_missing
        }


class FileParser:
    @staticmethod
    def parse_text(content: str) -> List[Card]:
        cards = []
        lines = content.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            card = FileParser._parse_line(line)
            if card:
                cards.append(card)
        return cards

    @staticmethod
    def _parse_line(line: str) -> Optional[Card]:
        # Printing-agnostic: "4 Lightning Bolt", "4x Solitude (MH2) #12 *F*",
        # "SB: 2 Force of Negation" all reduce to quantity + bare card name.
        # Set codes, collector numbers and foil markers are stripped.
        line = line.strip()
        if not line or line.startswith('#') or line.startswith('//'):
            return None
        line = re.sub(r'(?i)^SB:\s*', '', line)  # sideboard prefix first
        m = re.match(r'^(?:(\d+)\s*x?\s+)?(.*)$', line, re.IGNORECASE)
        if not m:
            return None
        qty = int(m.group(1)) if m.group(1) else 1
        name = m.group(2).strip()
        name = re.sub(r'(?i)\*F\*', '', name)  # foil marker
        has_set = bool(re.search(r'\([A-Za-z0-9]{2,6}\)', name))
        has_hash = '#' in name
        name = re.sub(r'\([A-Za-z0-9]{2,6}\)', '', name)  # set codes like (MH2)
        name = name.replace('#', ' ')
        if has_set or has_hash:
            # A trailing bare number after a set marker is a collector number,
            # not part of the name (names ending in digits without a set
            # marker are left untouched).
            name = re.sub(r'\s+\d+[a-z]?\s*$', '', name)
        name = re.sub(r'\s+', ' ', name).strip()
        if not name:
            return None
        return Card(name=name, quantity=qty)

    @staticmethod
    def parse_dek(content: str) -> List[Card]:
        cards = []
        in_sideboard = False
        for line in content.strip().split('\n'):
            line = line.strip()
            if not line or line.startswith('//'):
                continue
            if line.lower() == '[sideboard]' or line.lower() == 'sideboard':
                in_sideboard = True
                continue
            card = FileParser._parse_line(line)
            if card:
                cards.append(card)
        return cards

    @staticmethod
    def parse_csv(filepath: str) -> List[Card]:
        # Printing-agnostic: only Name + Quantity are used; Set/Number/Foil
        # columns are ignored so all printings merge into one entry.
        cards = []
        with open(filepath, 'r', encoding='utf-8-sig') as f:
            reader = csv.DictReader(f)
            for row in reader:
                name = row.get('Name') or row.get('Card') or row.get('Card Name') or ''
                if not name:
                    continue
                qty = int(row.get('Quantity') or row.get('Qty') or row.get('Count') or 1)
                cards.append(Card(name=name.strip(), quantity=qty))
        return cards

    @staticmethod
    def parse_xlsx(filepath: str) -> List[Card]:
        if not HAS_OPENPYXL:
            raise RuntimeError("openpyxl not installed. Run: pip install openpyxl")
        # Printing-agnostic: only Name + Quantity columns are used.
        cards = []
        wb = openpyxl.load_workbook(filepath)
        ws = wb.active
        headers = [cell.value for cell in ws[1]]
        name_idx = next((i for i, h in enumerate(headers) if h and 'name' in str(h).lower()), 0)
        qty_idx = next((i for i, h in enumerate(headers) if h and ('qty' in str(h).lower() or 'count' in str(h).lower())), 1)

        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[name_idx]:
                continue
            name = str(row[name_idx]).strip()
            qty = int(row[qty_idx]) if row[qty_idx] else 1
            cards.append(Card(name=name, quantity=qty))
        return cards

    @staticmethod
    def parse_docx(filepath: str) -> List[Card]:
        if not HAS_DOCX:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")
        cards = []
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            card = FileParser._parse_line(para.text)
            if card:
                cards.append(card)
        for table in doc.tables:
            for row in table.rows:
                text = " | ".join(cell.text for cell in row.cells)
                card = FileParser._parse_line(text)
                if card:
                    cards.append(card)
        return cards

    @staticmethod
    def parse_clipboard() -> List[Card]:
        if not HAS_CLIPBOARD:
            raise RuntimeError("pyperclip not installed. Run: pip install pyperclip")
        content = pyperclip.paste()
        return FileParser.parse_text(content)


class App:
    def __init__(self, plan_file: str, collection_file: str):
        self.plan = DeckPlan(plan_file)
        self.collection_file = collection_file
        self.collection = self.load_collection()
        self.tracker = ProgressTracker(self.plan, self.collection)

    def load_collection(self) -> Collection:
        if os.path.exists(self.collection_file):
            with open(self.collection_file, 'r') as f:
                return Collection.from_dict(json.load(f))
        return Collection()

    def save_collection(self):
        with open(self.collection_file, 'w') as f:
            json.dump(self.collection.to_dict(), f, indent=2)

    def import_file(self, filepath: str) -> int:
        ext = Path(filepath).suffix.lower()
        cards = []

        if ext in ('.txt', '.dek'):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            cards = FileParser.parse_dek(content) if ext == '.dek' else FileParser.parse_text(content)
        elif ext == '.csv':
            cards = FileParser.parse_csv(filepath)
        elif ext == '.xlsx':
            cards = FileParser.parse_xlsx(filepath)
        elif ext in ('.docx', '.doc'):
            cards = FileParser.parse_docx(filepath)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        for card in cards:
            self.collection.add(card)
        self.save_collection()
        return len(cards)

    def import_clipboard(self) -> int:
        cards = FileParser.parse_clipboard()
        for card in cards:
            self.collection.add(card)
        self.save_collection()
        return len(cards)

    def show_progress(self, phase_filter: str = None):
        overall = self.tracker.overall_progress()
        print(f"\n{'='*60}")
        print(f"OVERALL PROGRESS")
        print(f"{'='*60}")
        print(f"Unique Cards: {overall['unique_cards_owned']}/{overall['unique_cards_total']} ({overall['unique_percentage']}%)")
        print(f"Total Copies: {overall['copies_owned']}/{overall['copies_needed']} ({overall['copies_percentage']}%)")
        print(f"Estimated Value Missing: ${overall['estimated_value_missing']:,}")

        phases = self.tracker.all_phases_progress()
        if phase_filter:
            phases = [p for p in phases if phase_filter.lower() in p['phase'].lower()]

        print(f"\n{'='*60}")
        print(f"PHASE PROGRESS")
        print(f"{'='*60}")
        for p in phases:
            bar_len = 30
            filled = int(p['percentage'] / 100 * bar_len)
            bar = "#" * filled + "-" * (bar_len - filled)
            status = "[OK]" if p['percentage'] >= 100 else "[..]"
            print(f"\n{status} {p['phase']}")
            print(f"   [{bar}] {p['percentage']}% ({p['owned']}/{p['total']} cards)")
            print(f"   Missing: {p['missing_count']} cards | ${p['missing_value']:,} to complete")
            if p['missing']:
                for m in p['missing'][:5]:
                    print(f"     - {m['short']}x {m['name']} (${m['price']} ea) {m['notes']}")
                if len(p['missing']) > 5:
                    print(f"     ... and {len(p['missing']) - 5} more")

    def show_missing(self, phase_name: str = None):
        phases = self.tracker.all_phases_progress()
        if phase_name:
            phases = [p for p in phases if phase_name.lower() in p['phase'].lower()]

        for p in phases:
            if p['missing']:
                print(f"\n{p['phase']} - Missing ({p['missing_count']} cards, ${p['missing_value']:,}):")
                for m in p['missing']:
                    print(f"  {m['short']}x {m['name']} (${m['price']} ea) [{m['category']}] {m['notes']}")

    def show_collection(self, search: str = None):
        cards = list(self.collection.cards.values())
        if search:
            cards = [c for c in cards if search.lower() in c.name.lower()]
        cards.sort(key=lambda c: c.name.lower())
        print(f"\nCollection ({len(cards)} unique, {self.collection.total_cards()} total):")
        for c in cards:
            print(f"  {c.display_name()}")

    def interactive(self):
        print("\nMTG Modern Gateway Deck Tracker")
        print("Commands: progress, missing, collection, import, clipboard, save, quit")
        while True:
            try:
                cmd = input("\n> ").strip().lower()
                if cmd in ('q', 'quit', 'exit'):
                    break
                elif cmd in ('p', 'progress'):
                    self.show_progress()
                elif cmd.startswith('progress '):
                    self.show_progress(cmd.split(' ', 1)[1])
                elif cmd in ('m', 'missing'):
                    self.show_missing()
                elif cmd.startswith('missing '):
                    self.show_missing(cmd.split(' ', 1)[1])
                elif cmd in ('c', 'collection', 'list'):
                    self.show_collection()
                elif cmd.startswith('collection ') or cmd.startswith('list '):
                    self.show_collection(cmd.split(' ', 1)[1])
                elif cmd in ('i', 'import'):
                    path = input("File path: ").strip().strip('"')
                    if os.path.exists(path):
                        count = self.import_file(path)
                        print(f"Imported {count} cards from {path}")
                    else:
                        print("File not found")
                elif cmd in ('clip', 'clipboard', 'paste'):
                    try:
                        count = self.import_clipboard()
                        print(f"Imported {count} cards from clipboard")
                    except Exception as e:
                        print(f"Clipboard error: {e}")
                elif cmd == 'save':
                    self.save_collection()
                    print("Collection saved")
                elif cmd == 'help':
                    print("Commands: progress [filter], missing [filter], collection [search], import, clipboard, save, quit")
                else:
                    print("Unknown command. Type 'help' for commands.")
            except KeyboardInterrupt:
                break
            except EOFError:
                break
        print("\nGoodbye!")


def main():
    from paths import PLAN_FILE, COLLECTION_FILE
    plan_file, collection_file = PLAN_FILE, COLLECTION_FILE

    if not plan_file.exists():
        print(f"Plan file not found: {plan_file}")
        sys.exit(1)

    app = App(str(plan_file), str(collection_file))

    if len(sys.argv) > 1:
        cmd = sys.argv[1]
        if cmd == 'import' and len(sys.argv) > 2:
            count = app.import_file(sys.argv[2])
            print(f"Imported {count} cards")
        elif cmd == 'clipboard':
            count = app.import_clipboard()
            print(f"Imported {count} cards from clipboard")
        elif cmd == 'progress':
            filter_arg = sys.argv[2] if len(sys.argv) > 2 else None
            app.show_progress(filter_arg)
        elif cmd == 'missing':
            filter_arg = sys.argv[2] if len(sys.argv) > 2 else None
            app.show_missing(filter_arg)
        elif cmd == 'collection':
            search = sys.argv[2] if len(sys.argv) > 2 else None
            app.show_collection(search)
        else:
            print("Usage: python tracker.py [import <file>|clipboard|progress [filter]|missing [filter]|collection [search]]")
    else:
        app.interactive()


if __name__ == '__main__':
    main()